"""Generate polished research paper and executive summary."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

REPORTS_DIR = Path(__file__).resolve().parent.parent / "reports"


def add_bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_research_paper() -> None:
    doc = Document()
    title = doc.add_heading(
        "Workforce Attrition Patterns and Risk Hotspot Analysis at Palo Alto Networks",
        0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Unified Mentor Capstone Project | HR Analytics Research Paper")
    doc.add_paragraph("Author: [Your Name]")
    doc.add_paragraph("Date: September 2026")

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "Employee attrition in high-skill cybersecurity organizations can erode technical expertise, "
        "increase hiring costs, and disrupt strategic initiatives. This study analyzes 1,470 employee "
        "records from Palo Alto Networks to identify where attrition concentrates and which workforce "
        "segments face the highest exit risk. Using exploratory data analysis and a Random Forest "
        "classification model, the research establishes a 16.12% baseline attrition rate and highlights "
        "Sales, early-tenure employees, overtime workers, and single employees as primary risk hotspots. "
        "The findings support targeted, evidence-based retention strategies for HR leadership."
    )

    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "Palo Alto Networks operates in a fast-paced, knowledge-intensive cybersecurity environment where "
        "retaining skilled talent is strategically critical. Despite rich employee data, HR leadership "
        "lacked clear visibility into attrition patterns across departments, roles, demographics, tenure, "
        "and workload factors. This project delivers foundational diagnostic intelligence to shift "
        "workforce management from reactive to data-driven."
    )

    doc.add_heading("2. Problem Statement", level=1)
    add_bullet_list(
        doc,
        [
            "Which departments and job roles experience the highest attrition?",
            "Is attrition concentrated among specific age groups or tenure bands?",
            "Do overtime and business travel contribute to employee exits?",
            "Are early-career employees leaving more frequently than experienced staff?",
        ],
    )

    doc.add_heading("3. Dataset Description", level=1)
    doc.add_paragraph(
        "The dataset contains 1,470 employee records with 31 attributes including Age, Attrition (0/1), "
        "Department, JobRole, YearsAtCompany, OverTime, BusinessTravel, compensation metrics, satisfaction "
        "scores, and career progression indicators."
    )

    doc.add_heading("4. Methodology", level=1)
    doc.add_paragraph("The analysis followed a six-step diagnostic framework plus predictive modeling:")
    steps = [
        "Data Validation & Cleaning — attrition label normalization, duplicate removal, field standardization",
        "Overall Attrition Assessment — baseline turnover and retention proportions",
        "Department & Role Analysis — attrition rates and department-role heatmaps",
        "Demographic Analysis — age, gender, marital status, and education patterns",
        "Tenure & Career Stage Analysis — tenure bands and promotion stagnation impact",
        "Workload & Mobility Analysis — overtime, travel frequency, and commute distance",
        "Predictive Modeling — Random Forest classifier with balanced class weights",
    ]
    add_bullet_list(doc, steps)

    doc.add_heading("5. Results", level=1)

    doc.add_heading("5.1 Overall Attrition", level=2)
    doc.add_paragraph(
        "The organization-wide attrition rate is 16.12% (237 exits out of 1,470 employees), "
        "establishing the baseline for all comparative analyses."
    )

    doc.add_heading("5.2 Department & Role Hotspots", level=2)
    add_bullet_list(
        doc,
        [
            "Sales: 20.63% attrition (highest department)",
            "Human Resources: 19.05%",
            "Research & Development: 13.84%",
            "Sales Representative role: 39.76% (highest role)",
            "Laboratory Technician: 23.94%",
        ],
    )

    doc.add_heading("5.3 Demographic Patterns", level=2)
    add_bullet_list(
        doc,
        [
            "Age 18-25: 35.77% attrition — highest among age groups",
            "Single employees: 25.53% vs Married: 12.48%",
            "Male employees: 17.01% vs Female: 14.80%",
        ],
    )

    doc.add_heading("5.4 Tenure & Career Stage", level=2)
    add_bullet_list(
        doc,
        [
            "Early tenure (0-2 years): 29.82% attrition",
            "Mid tenure (3-6 years): 13.53%",
            "Senior tenure (7+ years): 10.68%",
            "Moderate promotion window (3-5 years since promotion): lowest attrition at 10.13%",
        ],
    )

    doc.add_heading("5.5 Workload & Mobility", level=2)
    add_bullet_list(
        doc,
        [
            "Overtime employees: 30.53% attrition vs 10.44% for non-overtime",
            "Frequent business travel: 24.91% vs Non-travel: 8.00%",
            "Employees living far from office (16+ miles): 20.67% attrition",
        ],
    )

    doc.add_heading("5.6 Predictive Model Performance", level=2)
    doc.add_paragraph(
        "A Random Forest model (200 trees, balanced class weights) achieved the following on a held-out test set:"
    )
    add_bullet_list(
        doc,
        [
            "Accuracy: 81.97%",
            "Precision: 44.83%",
            "Recall: 55.32%",
            "F1 Score: 49.52%",
            "ROC-AUC: 0.802",
        ],
    )
    doc.add_paragraph(
        "Top predictors: MonthlyIncome, Age, TotalWorkingYears, YearsAtCompany, and YearsWithCurrManager."
    )

    doc.add_heading("6. Recommendations", level=1)
    add_bullet_list(
        doc,
        [
            "Launch targeted retention programs in Sales and high-risk roles (Sales Representative, Laboratory Technician).",
            "Implement structured onboarding and mentorship for employees in their first two years.",
            "Review overtime policies and workload distribution in travel-intensive roles.",
            "Introduce promotion and career-path reviews for employees approaching stagnation.",
            "Deploy the Streamlit dashboard for ongoing HR monitoring and risk estimation.",
        ],
    )

    doc.add_heading("7. Conclusion", level=1)
    doc.add_paragraph(
        "This project provides foundational, diagnostic intelligence into employee attrition at Palo Alto Networks. "
        "By systematically analyzing attrition across departments, roles, demographics, tenure, and workload factors, "
        "HR leaders can move from intuition-based decisions to evidence-driven workforce management. "
        "The predictive model further enables proactive identification of at-risk employees before exit."
    )

    doc.save(REPORTS_DIR / "research_paper.docx")


def build_executive_summary() -> None:
    doc = Document()
    title = doc.add_heading("Executive Summary", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Workforce Attrition Analysis | Palo Alto Networks")
    doc.add_paragraph("Prepared for HR Leadership & Government Stakeholders")

    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(
        "This executive summary presents key findings from a comprehensive workforce attrition analysis "
        "designed to help Palo Alto Networks identify risk hotspots and prioritize retention investments."
    )

    doc.add_heading("Headline Metrics", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light List Accent 1"
    headers = table.rows[0].cells
    headers[0].text = "Metric"
    headers[1].text = "Value"
    metrics = [
        ("Total Employees Analyzed", "1,470"),
        ("Overall Attrition Rate", "16.12%"),
        ("Highest-Risk Department", "Sales (20.63%)"),
        ("Highest-Risk Role", "Sales Representative (39.76%)"),
        ("Early-Tenure Attrition (≤2 yrs)", "29.82%"),
        ("Overtime Attrition Rate", "30.53%"),
        ("Model ROC-AUC", "0.802"),
    ]
    for metric, value in metrics:
        row = table.add_row().cells
        row[0].text = metric
        row[1].text = value

    doc.add_heading("Critical Findings", level=1)
    add_bullet_list(
        doc,
        [
            "Attrition is not uniform — Sales and customer-facing roles face disproportionate turnover.",
            "The first two years represent a critical retention window with nearly 30% exit rate.",
            "Workload factors matter: overtime and frequent travel strongly correlate with attrition.",
            "Younger and single employees show elevated exit rates, suggesting lifestyle and career-stage factors.",
        ],
    )

    doc.add_heading("Recommended Actions", level=1)
    add_bullet_list(
        doc,
        [
            "Prioritize Sales department retention with role-specific interventions.",
            "Strengthen early-career support, mentoring, and clear progression pathways.",
            "Audit overtime and travel policies for high-risk teams.",
            "Use the live Streamlit dashboard for continuous monitoring and risk scoring.",
        ],
    )

    doc.add_heading("Deliverables", level=1)
    add_bullet_list(
        doc,
        [
            "Research paper with full EDA, insights, and recommendations",
            "Streamlit dashboard with interactive filters and predictive modeling",
            "Cleaned dataset and exported charts for reporting",
        ],
    )

    doc.save(REPORTS_DIR / "executive_summary.docx")


if __name__ == "__main__":
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    build_research_paper()
    build_executive_summary()
    print("Reports generated successfully.")
